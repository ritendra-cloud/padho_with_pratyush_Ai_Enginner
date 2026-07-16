import json
import time
import urllib.error
import urllib.request
from textwrap import shorten
from pathlib import Path
from pydantic import BaseModel


OLLAMA_URL = "http://localhost:11434/api/chat"
MODEL = "qwen2.5-coder:7b"


def ask_ollama(messages: list[dict[str, str]], schema: dict | None = None) -> str:
    payload = {
        "model": MODEL,
        "messages": messages,
        "stream": False,
        "options": {"temperature": 0},
    }
    if schema is not None:
        payload["format"] = schema

    request = urllib.request.Request(
        OLLAMA_URL,
        data=json.dumps(payload).encode("utf-8"),
        headers={"Content-Type": "application/json"},
        method="POST",
    )

    try:
        with urllib.request.urlopen(request, timeout=300) as response:
            return json.loads(response.read().decode("utf-8"))["message"]["content"]
    except urllib.error.URLError as error:
        raise RuntimeError(
            "Ollama is not reachable. Start it with `ollama serve` and try again."
        ) from error


job_description="""
Description
Do you want to solve real customer problems through innovative technology? Do you enjoy working on scalable services in a collaborative team environment? Do you want to see your code directly impact millions of customers worldwide?

At Amazon, we hire the best minds in technology to innovate and build on behalf of our customers. Customer obsession is part of our company DNA, which has made us one of the world's most beloved brands.

Our Software Development Engineers (SDEs) use modern technology to solve complex problems while seeing their work's impact first-hand. The challenges SDEs solve at Amazon are meaningful and influence millions of customers, sellers, and products globally. We seek individuals passionate about creating new products, features, and services while managing ambiguity in an environment where development cycles are measured in weeks, not years.

At Amazon, we believe in ownership at every level. As an SDE-I, you'll own the entire lifecycle of your code - from design through deployment and ongoing operations. This ownership mindset, combined with our commitment to operational excellence, ensures we deliver the highest quality solutions for our customers.

We're looking for curious minds who think big and want to define tomorrow's technology. At Amazon, you'll grow into the high-impact engineer you know you can be, supported by a culture of learning and mentorship. Every day brings exciting new challenges and opportunities for personal growth.
Key job responsibilities
• Collaborate and communicate effectively with experienced cross-disciplinary Amazonians to design, build, and operate innovative products and services that delight our customers, while participating in technical discussions to drive solutions forward.
• Design and develop scalable solutions using cloud-native architectures and microservices in a large distributed computing environment.
• Participate in code reviews and contribute to technical documentation.
• Build and maintain resilient distributed systems that are scalable, fault-tolerant, and cost-effective.
• Leverage and contribute to the development of GenAI and AI-powered tools to enhance development productivity while staying current with emerging technologies.
• Write clean, maintainable code following best practices and design patterns.
• Work in an agile environment practicing CI/CD principles while participating in operational responsibilities including on-call duties.
• Demonstrate operational excellence through monitoring, troubleshooting, and resolving production issues.
Basic Qualifications
- Experience with at least one general-purpose programming language such as Java, Python, C++, C#, Go, Rust, or TypeScript
- Experience with data structure implementation, basic algorithm development, and/or object-oriented design principles
- Currently has, or is in the process of obtaining a bachelor’s degree in Computer Science, Computer Engineering, Data Science, Information Systems, or related STEM fields
- Must be 18 years of age of older
Preferred Qualifications
- Experience from previous technical internship(s) or demonstrated project experience
- Experience with one or more of the following: AI tools for development productivity, Cloud platforms (preferably AWS), Database systems (SQL and NoSQL), Contributing to open-source projects, Version control systems, Debugging and troubleshooting complex systems
- Demonstrated ability to learn and adapt to new technologies quickly
- Basic understanding of software development lifecycle (SDLC)
- Strong problem-solving and analytical skills
- Excellent written and verbal communication skills
"""
class JobD(BaseModel):
    role: str
    required_skills: list[str]
    preferred_skills: list[str]
    minimum_experience: float | None
    education_requirements: list[str]
    responsibilities: list[str]

jobd_schema = JobD.model_json_schema()

system_prompt = f"""
You are an expert HR assistant.

Your job is to analyze job descriptions and extract
structured information from them.

Return ONLY valid JSON matching this schema:

{jobd_schema}

IMPORTANT:
Do NOT return the schema itself.
Do NOT return fields like "properties", "title" or "type".
Fill the schema with actual information extracted from the job description.

If minimum experience is not mentioned, return null.
If information for a list is missing, return an empty list.
Do not invent information.
"""

user_prompt = f"""
Analyze the following job description:

{job_description}
"""
message_system={
    "role" : "system",
    "content" : system_prompt
}
message_user={
    "role" : "user",
    "content" : user_prompt
}
messages=[message_system, message_user]

answer = ask_ollama(messages, schema=jobd_schema)

raw_json=answer
# print(raw_json)



job_data=json.loads(raw_json)

job = JobD(**job_data)

print(job.minimum_experience)
print(job.education_requirements)



#parse real
class MatchResult(BaseModel):
    score: float
    details: dict
class Experience(BaseModel):
    company: str | None = None
    role: str | None = None
    duration: str | None = None
    description: str | None = None
    skills_used: list[str] = []

class Resume(BaseModel):
    name: str | None = None
    email: str | None = None
    phone: str | None = None

    total_experience_years: float | None = None

    skills: list[str] = []
    experiences: list[Experience] = []
    education: list[str] = []
    projects: list[str] = []
    certifications: list[str] = []


resume_schema = Resume.model_json_schema()
def final_score(job,resume):
    match_schema = MatchResult.model_json_schema()
    prompt = f"""
    You are an HR recruiter.

    Compare the candidate's resume with the job description.

    JOB DESCRIPTION:
    {job.model_dump_json(indent=2)}

    CANDIDATE RESUME:
    {resume.model_dump_json(indent=2)}
    Return JSON matching this schema:

    {match_schema}

    Put these exact keys inside the details object:

    1. candidate_name
    2. matching_skills
    3. missing_important_skills
    4. experience_requirement_met
    5. overall_match_percentage
    6. final_verdict

    Keep the response concise and easy to read.
    """
    message={
        "role": "user",
        "content" : prompt
    }
    messages=[message]
    response = ask_ollama(messages, schema=match_schema)
    data = json.loads(response)
    return MatchResult(**data)
def parse_resume(resume_text):
    system_prompt = f"""
    You are an expert resume parser.

    Extract information from the resume based on its meaning,
    not only based on exact section headings.

    Different resumes may use different headings.

    For example:
    - Experience
    - Professional Experience
    - Work History
    - Employment
    - Internships

    These may all contain relevant experience.

    Skills may also appear in the skills section, work experience,
    internships or projects.

    Return ONLY valid JSON matching this schema:

    {resume_schema}

    Important rules:

    1. Do not invent information.
    2. If a value is not available, return null.
    3. If a list has no information, return an empty list.
    4. Include internships inside experiences.
    5. Extract skills mentioned across the entire resume.
    """
    user_prompt = f"""
    Parse the following resume:

    {resume_text}
    """
    message_system={
        "role" : "system",
        "content" : system_prompt
    }
    message_user={
        "role" : "user",
        "content" : user_prompt
    }
    messages=[message_system, message_user]
    raw_output = ask_ollama(messages, schema=resume_schema)
    data = json.loads(raw_output)
    resume = Resume(**data)
    return resume


from pypdf import PdfReader
from docx import Document

def read_pdf(file_path):
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

def read_docx(file_path):
    document = Document(file_path)
    text = ""
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"
    
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"
    return text


def read_resume(file_path):
    if file_path.suffix.lower() == ".pdf":
        return read_pdf(file_path)
    elif file_path.suffix.lower() == ".docx":
        return read_docx(file_path)
    else:
        return None


def make_short_text(value, width=35):
    if isinstance(value, list):
        value = ", ".join(str(item) for item in value)
    elif isinstance(value, bool):
        value = "Yes" if value else "No"
    elif value is None:
        value = ""
    else:
        value = str(value)
    return shorten(value, width=width, placeholder="...")


def print_results_table(results):
    headers = ["Candidate", "Score", "Matching Skills", "Missing Skills", "Exp Met", "Verdict"]
    rows = []

    for result in results:
        details = result["details"]
        rows.append([
            make_short_text(result["name"], 22),
            f"{result['score']}%",
            make_short_text(details.get("matching_skills"), 40),
            make_short_text(details.get("missing_important_skills"), 45),
            make_short_text(details.get("experience_requirement_met"), 8),
            make_short_text(details.get("final_verdict"), 55),
        ])

    widths = [
        max(len(str(row[index])) for row in [headers] + rows)
        for index in range(len(headers))
    ]

    def print_row(row):
        print("| " + " | ".join(str(cell).ljust(widths[index]) for index, cell in enumerate(row)) + " |")

    separator = "|-" + "-|-".join("-" * width for width in widths) + "-|"
    print("\nRESUME MATCH RESULTS")
    print_row(headers)
    print(separator)
    for row in rows:
        print_row(row)



# lets do it now
resume_folder = Path(__file__).parent / "resumes"
all_results=[]
for file_path in resume_folder.iterdir():
    #C:\Users\Pratyush\padho_with_pratyush\week1\day5\resumes\abhay resume new - Abhay Singh.pdf
    if file_path.suffix.lower() not in [".pdf", ".docx"]:
        continue
    print("\nProcessing:", file_path.name)
    resume_text = read_resume(file_path)
    parsed_resume=parse_resume(resume_text) # llm call1
    time.sleep(5)
    result = final_score(job, parsed_resume) #llm caLL2
    #score and details
    #acount chtgpt
    # request bhejna shhur krega millions
    #chattgot server jam ho jayega
    time.sleep(5)
    details = result.details
    candidate_name = details.get("candidate_name") or parsed_resume.name

    all_results.append({
        "name": candidate_name,
        "score": result.score,
        "details": details
    })

all_results.sort(
    key=lambda candidate: candidate["score"],
    reverse=True
)
print_results_table(all_results)
